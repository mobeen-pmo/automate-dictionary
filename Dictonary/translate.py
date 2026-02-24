"""
Automate Bilingual RPA Dictionary & Smart Translator (Pro Version)
Developed by Mirza Muhammad Mobeen
Refined for: Voice Commands (Speech-to-Text), Visual Progress, & Caching
"""

import streamlit as st
import pandas as pd
import os
import re
import io
import tempfile
import openpyxl
import time
import speech_recognition as sr  # NEW: For transcribing audio
from streamlit_mic_recorder import mic_recorder # NEW: For capturing audio
from deep_translator import GoogleTranslator
from docx import Document
from pdf2docx import Converter

# ──────────────────────────────────────────────
# 1. PAGE CONFIG
# ──────────────────────────────────────────────
st.set_page_config(
    page_title="Automate Keywords Dictionary",
    page_icon="🔤",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ──────────────────────────────────────────────
# 2. CUSTOM CSS
# ──────────────────────────────────────────────
st.markdown(
    """
    <style>
    /* ── Hide Streamlit defaults ── */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header[data-testid="stHeader"] {background: transparent !important;}
    header [data-testid="stToolbar"] {visibility: hidden;}

    /* Force sidebar toggle */
    button[data-testid="stSidebarCollapsedControl"] {
        visibility: visible !important;
        display: flex !important;
        color: #0072B5 !important;
        background: #FFFFFF !important;
        border: 1px solid #E2E8F0 !important;
        z-index: 99999 !important;
    }

    /* ── Fonts & General ── */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
    .block-container { padding-top: 2rem; padding-bottom: 5rem; }

    /* ── Typography ── */
    .hero-title {
        font-size: 2.6rem;
        font-weight: 800;
        background: linear-gradient(135deg, #0072B5 0%, #00A5E0 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0;
        letter-spacing: -0.5px;
    }
    .hero-subtitle {
        font-size: 1.05rem;
        color: #64748B !important;
        margin-top: 4px;
        margin-bottom: 2rem;
    }

    /* ── Inputs ── */
    div[data-testid="stTextInput"] > div > div > input,
    div[data-testid="stTextArea"] > div > div > textarea {
        font-size: 1.05rem;
        border: 2px solid #E2E8F0;
        border-radius: 14px;
        background: #FFFFFF !important;
        color: #1E293B !important;
    }
    div[data-testid="stTextInput"] > div > div > input:focus,
    div[data-testid="stTextArea"] > div > div > textarea:focus {
        border-color: #0072B5;
        box-shadow: 0 0 0 3px rgba(0, 114, 181, 0.12);
    }

    /* ── Term Card ── */
    .term-card {
        background: #FFFFFF !important;
        border: 1px solid #E2E8F0;
        border-radius: 16px;
        padding: 1.4rem 1.6rem;
        margin-bottom: 1rem;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .term-card:hover {
        transform: translateY(-3px);
        box-shadow: 0 8px 25px rgba(0, 114, 181, 0.10);
        border-color: #0072B5;
    }
    .cat-badge {
        display: inline-block;
        background: #EFF6FF !important;
        color: #0072B5 !important;
        padding: 3px 12px;
        border-radius: 8px;
        font-size: 0.73rem;
        font-weight: 600;
        margin-bottom: 1rem;
    }
    .cat-jp { color: #E11D48 !important; font-weight: 500; margin-left: 6px; }

    /* ── Grid Layouts ── */
    .card-grid {
        display: grid;
        grid-template-columns: 1fr auto 1fr;
        gap: 0.6rem;
        align-items: start;
    }
    .lang-block h4 {
        font-size: 0.7rem;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-bottom: 4px;
    }
    .lang-block .action-val { font-size: 1.05rem; font-weight: 700; color: #1E293B !important; }
    .lang-block .activity-val { font-size: 0.88rem; color: #64748B !important; }
    .en-block h4 { color: #0072B5 !important; }
    .jp-block h4 { color: #E11D48 !important; }
    .arrow-divider { font-size: 1.6rem; color: #CBD5E1 !important; padding-top: 1.2rem; }

    /* ── Sidebar & Footer ── */
    section[data-testid="stSidebar"] { background: #F8FAFC !important; border-right: 1px solid #E2E8F0; }
    .sidebar-title { font-size: 1.1rem; font-weight: 700; color: #0072B5 !important; margin-bottom: 0.3rem; }
    .sidebar-stat {
        background: #FFFFFF !important;
        border: 1px solid #E2E8F0;
        border-radius: 12px;
        padding: 1rem;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .sidebar-stat .stat-num { font-size: 1.8rem; font-weight: 800; color: #0072B5 !important; }
    .sidebar-stat .stat-label { font-size: 0.78rem; color: #94A3B8 !important; text-transform: uppercase; letter-spacing: 0.5px; }

    .custom-footer {
        position: fixed; bottom: 0; left: 0; width: 100%;
        background: #F8FAFC !important; border-top: 1px solid #E2E8F0;
        text-align: center; padding: 10px 0; font-size: 0.82rem; color: #94A3B8 !important; z-index: 9999;
    }
    .custom-footer span { color: #0072B5 !important; font-weight: 600; }
    
    /* ── Translation Highlight ── */
    .glossary-highlight {
        background-color: #FFF1F2;
        color: #E11D48;
        font-weight: 800;
        padding: 2px 6px;
        border-radius: 4px;
        border: 1px solid #FECDD3;
    }
    
    /* ── Result Box ── */
    .result-box {
        background: white;
        border: 2px solid #E2E8F0;
        border-radius: 14px;
        padding: 20px;
        min-height: 150px;
        font-size: 1.1rem;
        line-height: 1.6;
        color: #334155;
        margin-bottom: 10px;
    }
    
    /* ── Progress Bar Styling ── */
    .stProgress > div > div > div > div {
        background-color: #0072B5;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ──────────────────────────────────────────────
# 3. DATA LOADING & STATE INIT
# ──────────────────────────────────────────────
DATA_FILE = os.path.join(os.path.dirname(__file__), "Bilingual Automation Action and Activity Catalog.csv")

@st.cache_data
def load_data() -> pd.DataFrame:
    """Loads and cleans the CSV."""
    if not os.path.exists(DATA_FILE):
        return pd.DataFrame() 

    try:
        df = pd.read_csv(DATA_FILE, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(DATA_FILE, encoding="cp932")
    except Exception:
        return pd.DataFrame()

    if "Category" not in df.columns:
        if "Category (English)" in df.columns:
            df["Category"] = df["Category (English)"]
        elif "Category (Japanese)" in df.columns:
            df["Category"] = df["Category (Japanese)"]
    
    if "Category (Japanese)" not in df.columns:
        df["Category (Japanese)"] = ""

    required_cols = [
        "Category", "Category (Japanese)", 
        "Action (English)", "Action (Japanese)", 
        "Activity (English)", "Activity (Japanese)"
    ]
    for col in required_cols:
        if col not in df.columns:
            df[col] = ""

    if "Source" in df.columns:
        df = df.drop(columns=["Source"])
        
    df = df.drop_duplicates()
    
    for col in df.columns:
        df[col] = df[col].astype(str).str.strip().replace("nan", "")
        
    return df

df = load_data()

# INITIALIZE SESSION STATE FOR TEXT INPUT
if 'trans_input' not in st.session_state:
    st.session_state.trans_input = ""

if df.empty:
    st.error(f"⚠️ **Error:** Data file not found. Ensure `{os.path.basename(DATA_FILE)}` is in the directory.")
    st.stop()

# ──────────────────────────────────────────────
# 4. SIDEBAR
# ──────────────────────────────────────────────
with st.sidebar:
    st.markdown('<p class="sidebar-title">🔤 Automate Dictionary</p>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:0.82rem;color:#94A3B8;margin-top:-4px;">Bilingual RPA Dictionary</p>', unsafe_allow_html=True)
    st.markdown("---")

    cat_map = {} 
    for _, row in df[["Category", "Category (Japanese)"]].drop_duplicates().iterrows():
        en = row["Category"]
        jp = row["Category (Japanese)"]
        label = f"{en} ({jp})" if jp and jp != en and jp != "" else en
        cat_map[label] = en

    category_labels = sorted(cat_map.keys())
    selected_labels = st.multiselect(
        "📂 Filter by Category",
        options=category_labels,
        default=[],
        help="Leave empty to search all categories.",
    )
    selected_cats = [cat_map[lbl] for lbl in selected_labels]

    st.markdown("---")
    col_s1, col_s2 = st.columns(2)
    with col_s1:
        st.markdown(f'<div class="sidebar-stat"><div class="stat-num">{len(df)}</div><div class="stat-label">Terms</div></div>', unsafe_allow_html=True)
    with col_s2:
        st.markdown(f'<div class="sidebar-stat"><div class="stat-num">{df["Category"].nunique()}</div><div class="stat-label">Cats</div></div>', unsafe_allow_html=True)

    st.markdown("---")
    view_all = st.toggle("📊 View All Data", value=False)
    st.markdown("---")
    st.markdown('<p style="text-align:center;font-size:0.75rem;color:#CBD5E1;">Developed by<br><b>Mirza Muhammad Mobeen</b></p>', unsafe_allow_html=True)


# ──────────────────────────────────────────────
# 5. TRANSLATION LOGIC (TEXT & DOCS)
# ──────────────────────────────────────────────

def smart_translate_text(text, glossary_df, direction="En_to_Jp", return_html=True, cache=None):
    if not isinstance(text, str) or not text.strip():
        return text

    if cache is not None and text in cache:
        return cache[text]

    if direction == "En_to_Jp":
        src_lang, tgt_lang = 'en', 'ja'
        action_map = dict(zip(glossary_df["Action (English)"].str.lower(), glossary_df["Action (Japanese)"]))
        activity_map = dict(zip(glossary_df["Activity (English)"].str.lower(), glossary_df["Activity (Japanese)"]))
    else:
        src_lang, tgt_lang = 'ja', 'en'
        action_map = dict(zip(glossary_df["Action (Japanese)"], glossary_df["Action (English)"]))
        activity_map = dict(zip(glossary_df["Activity (Japanese)"], glossary_df["Activity (English)"]))
    
    full_map = {**action_map, **activity_map}

    pattern = re.compile(r'("([^"]+)")|(\'([^\']+)\')|(「([^」]+)」)|(『([^』]+)』)')
    placeholders = {}
    
    def replacer(match):
        if match.group(2): term = match.group(2)
        elif match.group(4): term = match.group(4)
        elif match.group(6): term = match.group(6)
        elif match.group(8): term = match.group(8)
        else: return match.group(0)

        lookup_key = term.lower() if direction == "En_to_Jp" else term
        
        if lookup_key in full_map:
            target_term = full_map[lookup_key]
            key = f"[ID{len(placeholders)}]" 
            placeholders[key] = target_term
            return key
        else:
            return match.group(0)

    processed_text = pattern.sub(replacer, text)
    
    try:
        translator = GoogleTranslator(source=src_lang, target=tgt_lang)
        translated_text = translator.translate(processed_text)
    except Exception as e:
        return f"Error: {str(e)}"

    if not translated_text:
        return text

    final_text = translated_text 
    for key, term in placeholders.items():
        if direction == "En_to_Jp":
            if return_html:
                formatted_term = f"「<span class='glossary-highlight'>{term}</span>」"
            else:
                formatted_term = f"「{term}」"
        else:
            if return_html:
                formatted_term = f""" "<span class='glossary-highlight'>{term}</span>" """
            else:
                formatted_term = f'"{term}"'
        
        escaped_key_regex = re.escape(key).replace(r"\[", r"\[\s*").replace(r"\]", r"\s*\]")
        final_text = re.sub(escaped_key_regex, formatted_term, final_text)

    if cache is not None:
        cache[text] = final_text

    return final_text

# ──────────────────────────────────────────────
# 6. DOCUMENT PROCESSING
# ──────────────────────────────────────────────

def translate_docx_file(input_file, glossary_df, direction, progress_bar=None, status_text=None):
    doc = Document(input_file)
    total_items = len(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            total_items += len(row.cells)
    
    if total_items == 0: total_items = 1
    current_item = 0
    translation_memory = {} 

    def update_prog():
        nonlocal current_item
        current_item += 1
        pct = int((current_item / total_items) * 100)
        if progress_bar: progress_bar.progress(min(current_item / total_items, 1.0))
        if status_text: status_text.text(f"Processing... {pct}%")

    for para in doc.paragraphs:
        if para.text.strip():
            translated = smart_translate_text(para.text, glossary_df, direction, return_html=False, cache=translation_memory)
            para.text = translated
        update_prog()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text = ""
                for p in cell.paragraphs:
                    full_text += p.text + "\n"
                if full_text.strip():
                    translated = smart_translate_text(full_text.strip(), glossary_df, direction, return_html=False, cache=translation_memory)
                    cell.text = translated
                update_prog()
    
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer

def convert_and_translate_pdf(input_file, glossary_df, direction, progress_bar=None, status_text=None):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tf_input:
        tf_input.write(input_file.read())
        temp_input_path = tf_input.name
    temp_docx_path = temp_input_path.replace(".pdf", ".docx")
    try:
        if status_text: status_text.text("Converting PDF to editable format...")
        cv = Converter(temp_input_path)
        cv.convert(temp_docx_path, start=0, end=None)
        cv.close()
        
        if status_text: status_text.text("Starting Translation...")
        with open(temp_docx_path, "rb") as f:
            docx_buffer = translate_docx_file(f, glossary_df, direction, progress_bar, status_text)
        return docx_buffer
    finally:
        if os.path.exists(temp_input_path): os.remove(temp_input_path)
        if os.path.exists(temp_docx_path): os.remove(temp_docx_path)

def translate_excel_file(input_file, glossary_df, direction, is_legacy=False, progress_bar=None, status_text=None):
    output_buffer = io.BytesIO()
    translation_memory = {}
    if is_legacy:
        try:
            xls = pd.read_excel(input_file, sheet_name=None)
            total_cells = sum(df.size for df in xls.values()) 
            if total_cells == 0: total_cells = 1
            current_count = 0
            with pd.ExcelWriter(output_buffer, engine='openpyxl') as writer:
                for sheet_name, df_sheet in xls.items():
                    def progress_map(x):
                        nonlocal current_count
                        current_count += 1
                        if current_count % 10 == 0: 
                             pct = int((current_count / total_cells) * 100)
                             if progress_bar: progress_bar.progress(min(current_count / total_cells, 1.0))
                             if status_text: status_text.text(f"Processing... {pct}%")
                        return smart_translate_text(x, glossary_df, direction, False, cache=translation_memory) if isinstance(x, str) else x

                    df_sheet = df_sheet.map(progress_map)
                    df_sheet.columns = [smart_translate_text(str(c), glossary_df, direction, False, cache=translation_memory) for c in df_sheet.columns]
                    df_sheet.to_excel(writer, sheet_name=sheet_name, index=False)
            output_buffer.seek(0)
            return output_buffer
        except Exception as e:
            return None 
    else:
        wb = openpyxl.load_workbook(input_file)
        cells_to_process = []
        if status_text: status_text.text("Analyzing file structure...")
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and cell.value.strip():
                        cells_to_process.append(cell)
        
        total_cells = len(cells_to_process)
        if total_cells == 0: total_cells = 1
        
        for i, cell in enumerate(cells_to_process):
            translated = smart_translate_text(cell.value, glossary_df, direction, return_html=False, cache=translation_memory)
            cell.value = translated
            if i % 5 == 0 or i == total_cells - 1:
                pct = int(((i + 1) / total_cells) * 100)
                if progress_bar: progress_bar.progress((i + 1) / total_cells)
                if status_text: status_text.text(f"Processing... {pct}%")
        wb.save(output_buffer)
        output_buffer.seek(0)
        return output_buffer

def translate_csv_file(input_file, glossary_df, direction, progress_bar=None, status_text=None):
    try:
        df_csv = pd.read_csv(input_file)
        translation_memory = {}
        total_cells = df_csv.size
        if total_cells == 0: total_cells = 1
        current_count = 0
        def progress_map(x):
            nonlocal current_count
            current_count += 1
            if current_count % 20 == 0:
                 pct = int((current_count / total_cells) * 100)
                 if progress_bar: progress_bar.progress(min(current_count / total_cells, 1.0))
                 if status_text: status_text.text(f"Processing... {pct}%")
            return smart_translate_text(x, glossary_df, direction, False, cache=translation_memory) if isinstance(x, str) else x

        df_csv = df_csv.map(progress_map)
        df_csv.columns = [smart_translate_text(str(c), glossary_df, direction, False, cache=translation_memory) for c in df_csv.columns]
        if progress_bar: progress_bar.progress(1.0)
        if status_text: status_text.text("Processing... 100%")
        output_buffer = io.BytesIO()
        df_csv.to_csv(output_buffer, index=False, encoding='utf-8-sig')
        output_buffer.seek(0)
        return output_buffer
    except Exception as e:
        return None

# ──────────────────────────────────────────────
# 7. MAIN PAGE & TABS
# ──────────────────────────────────────────────
st.markdown('<p class="hero-title">Automate Keywords Dictionary</p>', unsafe_allow_html=True)
st.markdown('<p class="hero-subtitle">Bilingual Reference & Smart Translator</p>', unsafe_allow_html=True)

# TABS
tab_dict, tab_trans, tab_files = st.tabs(["📖 Dictionary Search", "🤖 Smart Translator", "📄 Document Translator"])

# ──────────────────────────────────────────────
# TAB 1: DICTIONARY SEARCH
# ──────────────────────────────────────────────
with tab_dict:
    query = st.text_input("search_input", placeholder="Search term (e.g., 'Excel', 'Browser')...", label_visibility="collapsed")
    filtered = df.copy()
    if selected_cats:
        filtered = filtered[filtered["Category"].isin(selected_cats)]

    if query.strip():
        q = query.strip().lower()
        cols = ["Category", "Category (Japanese)", "Action (English)", "Action (Japanese)", "Activity (English)", "Activity (Japanese)"]
        mask = pd.Series(False, index=filtered.index)
        for col in cols:
            mask = mask | filtered[col].str.lower().str.contains(q, na=False)
        filtered = filtered[mask]

    st.markdown("---")
    
    if view_all:
        st.dataframe(filtered, use_container_width=True, hide_index=True)
    else:
        if filtered.empty:
            st.markdown('<div class="no-results"><div class="emoji">🔍</div><p>No terms found.</p></div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<span class="cat-badge" style="margin-bottom:15px;">Found {len(filtered)} terms</span>', unsafe_allow_html=True)
            display_df = filtered.head(100)
            for _, row in display_df.iterrows():
                cat = row.get("Category", "")
                cat_jp = row.get("Category (Japanese)", "")
                en_act = row.get("Action (English)", "")
                jp_act = row.get("Action (Japanese)", "")
                en_activity = row.get("Activity (English)", "")
                jp_activity = row.get("Activity (Japanese)", "")
                cat_display = f'{cat} <span class="cat-jp">({cat_jp})</span>' if cat_jp and cat_jp != cat else cat

                card_html = f"""
                <div class="term-card">
                    <div class="cat-badge">{cat_display}</div>
                    <div class="card-grid">
                        <div class="lang-block en-block">
                            <h4>🇬🇧 English</h4>
                            <div class="action-val">{en_act}</div>
                            <div class="activity-val">{en_activity}</div>
                        </div>
                        <div class="arrow-divider">→</div>
                        <div class="lang-block jp-block">
                            <h4>🇯🇵 Japanese</h4>
                            <div class="action-val">{jp_act}</div>
                            <div class="activity-val">{jp_activity}</div>
                        </div>
                    </div>
                </div>
                """
                st.markdown(card_html, unsafe_allow_html=True)

# ──────────────────────────────────────────────
# TAB 2: SMART TRANSLATOR (VOICE ENABLED)
# ──────────────────────────────────────────────
with tab_trans:
    st.markdown("""
    <div style="background:#F0F9FF;padding:15px;border-radius:10px;border:1px solid #BAE6FD;margin-bottom:20px;">
        <strong style="color:#0072B5">🤖 Smart Detection:</strong><br>
        Put Automate terms in quotes. The app will detect them, insert the official dictionary term, 
        and format the sentence professionally.<br>
        <small style="color:#64748B">The translator preserves the formality of your input.</small>
    </div>
    """, unsafe_allow_html=True)

    direction_mode = st.radio(
        "Translation Direction (Text):",
        ["🇺🇸 English ➝ 🇯🇵 Japanese", "🇯🇵 Japanese ➝ 🇺🇸 English"],
        horizontal=True,
        key="text_dir"
    )
    
    # Define Lang Codes for Speech Rec
    if "English" in direction_mode.split("➝")[0]:
        dir_code = "En_to_Jp"
        src_label = "🇺🇸 English Input"
        tgt_label = "🇯🇵 Japanese Output"
        placeholder_txt = 'Example: Dear Team, I want to use the "Excel Open" action for the process.'
        speech_lang = "en-US"
    else:
        dir_code = "Jp_to_En"
        src_label = "🇯🇵 Japanese Input"
        tgt_label = "🇺🇸 English Output"
        placeholder_txt = 'Example: 「Excel 開く」アクションを使用したいです。'
        speech_lang = "ja-JP"

    col_t1, col_t2 = st.columns(2)
    
    with col_t1:
        st.subheader(src_label)
        
        # 🎙️ AUDIO RECORDER LOGIC
        col_mic, col_label = st.columns([1, 4])
        with col_mic:
            # Mic recorder captures bytes
            audio = mic_recorder(
                start_prompt="🎤 Speak",
                stop_prompt="🛑 Stop",
                just_once=True,
                key='recorder'
            )
        
        # PROCESS AUDIO
        if audio:
            with st.spinner("🎧 Transcribing audio..."):
                try:
                    # Save bytes to temp wav file
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_wav:
                        temp_wav.write(audio['bytes'])
                        temp_wav_path = temp_wav.name
                    
                    # Recognize speech
                    r = sr.Recognizer()
                    with sr.AudioFile(temp_wav_path) as source:
                        audio_data = r.record(source)
                        # Detect text
                        text = r.recognize_google(audio_data, language=speech_lang)
                        # Update Session State
                        st.session_state.trans_input = text
                    
                    # Cleanup
                    os.remove(temp_wav_path)
                    st.rerun() # Refresh to show text in box
                except sr.UnknownValueError:
                    st.warning("Could not understand audio.")
                except sr.RequestError as e:
                    st.error(f"Speech Service Error: {e}")
                except Exception as e:
                    st.error(f"Error: {e}")

        # TEXT AREA (Linked to Session State)
        def update_text_input():
            st.session_state.trans_input = st.session_state.widget_input

        source_text = st.text_area(
            "Write here...", 
            height=200, 
            placeholder=placeholder_txt,
            key="widget_input",
            value=st.session_state.trans_input,
            on_change=update_text_input
        )
        
        btn = st.button("Translate Text", type="primary", use_container_width=True)

    with col_t2:
        st.subheader(tgt_label)
        
        if btn and source_text:
            with st.spinner("Translating..."):
                html_result = smart_translate_text(source_text, df, direction=dir_code, return_html=True)
                plain_result = smart_translate_text(source_text, df, direction=dir_code, return_html=False)
            
            st.markdown(f'<div class="result-box">{html_result}</div>', unsafe_allow_html=True)
            st.caption("📋 Copy raw text:")
            st.code(plain_result, language=None)
            
        elif not source_text and btn:
            st.warning("Please enter text to translate.")
        else:
            st.markdown('<div class="result-box" style="color:#94A3B8;display:flex;align-items:center;justify-content:center;">Translation will appear here...</div>', unsafe_allow_html=True)

# ──────────────────────────────────────────────
# TAB 3: DOCUMENT TRANSLATOR
# ──────────────────────────────────────────────
with tab_files:
    st.markdown("""
    <div style="background:#FDF2F8;padding:15px;border-radius:10px;border:1px solid #FCC2D7;margin-bottom:20px;">
        <strong style="color:#BE185D">📄 File Translator (Fast & Visual):</strong><br>
        Upload Docs or Spreadsheets. The system will translate while preserving layout.
        <br>
        <small>Includes <b>Smart Caching</b>: Repeated words are translated instantly to save time.</small>
    </div>
    """, unsafe_allow_html=True)

    file_dir_mode = st.radio(
        "Translation Direction (File):",
        ["🇺🇸 English ➝ 🇯🇵 Japanese", "🇯🇵 Japanese ➝ 🇺🇸 English"],
        horizontal=True,
        key="file_dir"
    )
    
    f_dir_code = "En_to_Jp" if "English" in file_dir_mode.split("➝")[0] else "Jp_to_En"

    uploaded_file = st.file_uploader("Upload your document", type=["docx", "pdf", "xlsx", "xls", "csv"])

    if uploaded_file is not None:
        file_ext = uploaded_file.name.split(".")[-1].lower()
        st.info(f"File detected: {uploaded_file.name}")
        
        if st.button("🚀 Start Translation", type="primary"):
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                output_data = None
                out_name = f"Translated_{uploaded_file.name}"
                mime_type = "application/octet-stream"
                
                if file_ext == "docx":
                    output_data = translate_docx_file(uploaded_file, df, f_dir_code, progress_bar, status_text)
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif file_ext == "pdf":
                    output_data = convert_and_translate_pdf(uploaded_file, df, f_dir_code, progress_bar, status_text)
                    out_name = out_name.replace(".pdf", ".docx")
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif file_ext == "xlsx":
                    output_data = translate_excel_file(uploaded_file, df, f_dir_code, is_legacy=False, progress_bar=progress_bar, status_text=status_text)
                    mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                elif file_ext == "xls":
                    output_data = translate_excel_file(uploaded_file, df, f_dir_code, is_legacy=True, progress_bar=progress_bar, status_text=status_text)
                    out_name = out_name.replace(".xls", ".xlsx")
                    mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                elif file_ext == "csv":
                    output_data = translate_csv_file(uploaded_file, df, f_dir_code, progress_bar, status_text)
                    mime_type = "text/csv"
                
                if output_data:
                    progress_bar.progress(100)
                    status_text.success("✅ Translation Complete!")
                    st.download_button(
                        label="📥 Download Translated Document",
                        data=output_data,
                        file_name=out_name,
                        mime=mime_type
                    )
                else:
                    status_text.error("Failed to generate output data.")
                    
            except Exception as e:
                status_text.error(f"Error processing file: {str(e)}")

# ──────────────────────────────────────────────
# 8.FOOTER
# ──────────────────────────────────────────────
st.markdown(
    '<div class="custom-footer">© 2026 | Developed by <span>Mirza Muhammad Mobeen</span></div>',
    unsafe_allow_html=True,
)
